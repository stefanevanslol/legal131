import sys
import os
from docx import Document

# Add project root to path
sys.path.append(os.getcwd())

from backend.core.document_generator import DocumentGenerator

TEMPLATE_PATH = r"backend/templates/Natale - GEICO BI Demand (3) (6).docx"
OUTPUT_DIR = "backend/outputs"

def verify_render():
    if not os.path.exists(TEMPLATE_PATH):
        print(f"CRITICAL: Template not found at {TEMPLATE_PATH}")
        return

    doc_gen = DocumentGenerator(TEMPLATE_PATH)
    
    # Test Data 1: Input uses CORRECT spelling "Diagnoses"
    # Expected: The logic should map this to "Diagnosses" if the template needs it.
    data_1 = {
        "clientName": "TEST_CLIENT_1",
        "TreatmentsAndDiagnoses": "RENDERED_FROM_CORRECT_KEY",
        "TreatmentsAndDiagnosses2": "RENDERED_FROM_CORRECT_KEY_2" 
    }
    
    # Test Data 2: Input uses TYPO spelling "Diagnosses"
    # Expected: The logic should just pass it through (or map to Diagnoses).
    data_2 = {
        "clientName": "TEST_CLIENT_2",
        "TreatmentsAndDiagnosses": "RENDERED_FROM_TYPO_KEY",
        "TreatmentsAndDiagnosses2": "RENDERED_FROM_TYPO_KEY_2"
    }

    out_1 = os.path.join(OUTPUT_DIR, "verify_render_correct.docx")
    out_2 = os.path.join(OUTPUT_DIR, "verify_render_typo.docx")

    print(f"Generating {out_1}...")
    doc_gen.generate_demand_letter(data_1, out_1)
    
    print(f"Generating {out_2}...")
    doc_gen.generate_demand_letter(data_2, out_2)

    # Verify Content
    check_file(out_1, "RENDERED_FROM_CORRECT_KEY")
    check_file(out_2, "RENDERED_FROM_TYPO_KEY")

def check_file(path, expected_text):
    if not os.path.exists(path):
        print(f"FAIL: File {path} was not created.")
        return

    doc = Document(path)
    full_text = []
    for p in doc.paragraphs:
        full_text.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    full_text.append(p.text)
    
    content = "\n".join(full_text)
    
    if expected_text in content:
        print(f"PASS: '{expected_text}' found in {path}")
    else:
        print(f"FAIL: '{expected_text}' NOT found in {path}")
        # print(f"DEBUG Snippet: {content[:500]}...")

if __name__ == "__main__":
    verify_render()
