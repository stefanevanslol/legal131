import sys
import os
from docx import Document

# Add project root to path
sys.path.append(os.getcwd())

from backend.core.document_generator import DocumentGenerator

TEMPLATE_PATH = r"backend/templates/Natale - GEICO BI Demand (3) (6).docx"
OUTPUT_DIR = "backend/outputs"

def verify_richtext():
    if not os.path.exists(TEMPLATE_PATH):
        print(f"CRITICAL: Template not found at {TEMPLATE_PATH}")
        return

    doc_gen = DocumentGenerator(TEMPLATE_PATH)
    
    # Test Data: Use Markdown bolding to trigger RichText conversion
    data = {
        "clientName": "RICHTEXT_TEST_CLIENT",
        # This will become a RichText object in the generator
        "TreatmentsAndDiagnosses": "Start of line **BOLD TEXT** end of line.", 
        "TreatmentsAndDiagnosses2": "Normal text." 
    }
    
    out_path = os.path.join(OUTPUT_DIR, "verify_richtext.docx")

    print(f"Generating {out_path} with RichText input...")
    doc_gen.generate_demand_letter(data, out_path)
    
    # Verify Content
    check_file(out_path, "BOLD TEXT")

def check_file(path, expected_text):
    if not os.path.exists(path):
        print(f"FAIL: File {path} was not created.")
        return

    doc = Document(path)
    full_text = []
    for p in doc.paragraphs:
        full_text.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    full_text.append(p.text)
    
    content = "\n".join(full_text)
    
    # Note: docx.Document(path).paragraphs[...].text DOES contain the text even if it's bold.
    if expected_text in content:
        print(f"PASS: Found '{expected_text}' in document. RichText rendering works.")
        # Optional: Print surrounding context
        idx = content.find(expected_text)
        start = max(0, idx - 20)
        end = min(len(content), idx + 20)
        print(f"Context: ...{content[start:end]}...")
        
    else:
        print(f"FAIL: '{expected_text}' NOT found in document. RichText rendering FAILED.")

if __name__ == "__main__":
    verify_richtext()
