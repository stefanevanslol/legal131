from docx import Document
import re
import os

# Path to the user's file
template_path = r"c:\Users\max'\Downloads\Natale - GEICO BI Demand (3) (5).docx"

def extract_variables():
    if not os.path.exists(template_path):
        print(f"Error: File not found at {template_path}")
        return

    try:
        doc = Document(template_path)
    except Exception as e:
        print(f"Error reading docx: {e}")
        return

    text_content = []
    
    # Collect all text from paragraphs
    for p in doc.paragraphs:
        text_content.append(p.text)
        
    # Collect all text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text_content.append(p.text)

    full_text = "\n".join(text_content)
    
    # Regex patterns
    curly_pattern = r"\{\{([^\}]+)\}\}"  # Matches {{value}}
    bracket_pattern = r"\[([^\]]+)\]"    # Matches [value]
    
    # Find matches
    curly_vars = list(set(re.findall(curly_pattern, full_text)))
    bracket_vars = list(set(re.findall(bracket_pattern, full_text)))
    
    print("--- FOUND VARIABLES ---")
    if curly_vars:
        print("\nCurly Brace Variables {{x}}:")
        for v in sorted(curly_vars):
            print(f"- {v}")
            
    if bracket_vars:
        print("\nSquare Bracket Variables [x]:")
        for v in sorted(bracket_vars):
            print(f"- {v}")

    # Fallback: Look for suspicious ALL CAPS keys if few markers are found
    if len(curly_vars) + len(bracket_vars) < 3:
        print("\nPossible ALL CAPS Placeholders (Scanning...):")
        # Heuristic: Words in all caps, possibly with underscores, ignoring common short words
        caps_pattern = r"\b[A-Z0-9_]{3,}(?: [A-Z0-9_]{3,})*\b"
        all_caps = set(re.findall(caps_pattern, full_text))
        
        # Filter out likely noise
        ignore_list = {"THE", "AND", "FOR", "WITH", "FROM", "DATE", "PAGE", "OF", "DEMAND", "LETTER", "REGARDING", "CLAIM", "NUMBER"}
        potential_caps = [w for w in all_caps if w not in ignore_list and len(w) > 3]
        
        for v in sorted(potential_caps):
            print(f"- {v}")

if __name__ == "__main__":
    extract_variables()
