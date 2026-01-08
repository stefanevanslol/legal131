from docx import Document
import re
import os
import json

# Path to the new template
template_path = r"backend/templates/Natale - GEICO BI Demand (3) (5).docx"

def extract_context():
    if not os.path.exists(template_path):
        print(f"Error: File not found at {template_path}")
        return

    doc = Document(template_path)
    
    # Map of variable -> surrounding context
    context_map = {}
    
    # Regex for {{variable}}
    pattern = re.compile(r"\{\{([^\}]+)\}\}")
    
    for paragraph in doc.paragraphs:
        matches = pattern.findall(paragraph.text)
        for var in matches:
            # Grab the whole paragraph as context
            context_map[var] = paragraph.text.strip()

    # Also check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                matches = pattern.findall(text)
                for var in matches:
                    context_map[var] = text

    print(json.dumps(context_map, indent=2))

if __name__ == "__main__":
    extract_context()
