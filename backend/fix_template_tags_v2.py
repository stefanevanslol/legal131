from docx import Document
import os

# Target the specific file mentioned in earlier steps
template_path = r"backend/templates/Natale - GEICO BI Demand (3) (6).docx"

def fix_tags():
    if not os.path.exists(template_path):
        print(f"File not found: {template_path}")
        return

    print(f"Opening {template_path}...")
    doc = Document(template_path)
    
    replacements = {
        "{{r TreatmentsAndDiagnosses}}": "{{TreatmentsAndDiagnosses}}",
        "{{r TreatmentsAndDiagnosses2}}": "{{TreatmentsAndDiagnosses2}}",
        "{{r complaintsByclient}}": "{{complaintsByclient}}"
    }

    count = 0
    
    # helper to replace in text
    def replace_in_text(text):
        nonlocal count
        original = text
        for old, new in replacements.items():
            if old in text:
                text = text.replace(old, new)
                count += 1
                print(f"Replaced: {old} -> {new}")
        return text

    # Iterate paragraphs
    for p in doc.paragraphs:
        if "{{" in p.text:
            # We need to be careful. python-docx text is read-only on the paragraph, 
            # we have to modify runs. But often tags are split across runs. 
            # For this specific case, if the tag is simple, we might get lucky.
            # If not, we might need a more robust approach (clearing runs and resetting text).
            # Let's try the simple replace first, if it fails (because text is split), checking logs will tell us.
            # Actually, check_docx_tags found the EXACT string "{{r ...}}", so they are likely contiguous text.
            
            # However, p.text setter replaces all runs with a single run preserving style of first run? No, it often loses formatting.
            # A safer way requires iterating runs, but that is complex if split.
            # Given the previous context showing "r " prefix, it's likely a manual typo or a specific find/replace artifact.
            
            # Strategy: Check if the exact string exists in p.text. If so, try to replace in runs if possible, 
            # otherwise simplistic replacement might be safer if we accept some formatting risk, but user said "dont change anything else".
            # Let's try to find the run containing the text first.
            
            full_text = p.text
            for old, new in replacements.items():
                if old in full_text:
                    # It's in the paragraph. Let's see if it's in a single run.
                    replaced_in_run = False
                    for run in p.runs:
                        if old in run.text:
                            run.text = run.text.replace(old, new)
                            count += 1
                            print(f"Replaced in run: {old} -> {new}")
                            replaced_in_run = True
                    
                    if not replaced_in_run:
                        # It's split across runs. This is the hard case. 
                        # To avoid breaking formatting, we should probably warn or try a more aggressive fix used in 'fix_template_tags.py' style.
                        # But let's see if we can just do p.text assignment if it's a simple placeholder line.
                        print(f"WARNING: Found {old} but it is split across runs. Attemping text replacement (may reset local formatting).")
                        p.text = p.text.replace(old, new)
                        count += 1

    # Iterate tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if "{{" in p.text:
                        full_text = p.text
                        for old, new in replacements.items():
                            if old in full_text:
                                replaced_in_run = False
                                for run in p.runs:
                                    if old in run.text:
                                        run.text = run.text.replace(old, new)
                                        count += 1
                                        print(f"Replaced in table run: {old} -> {new}")
                                        replaced_in_run = True
                                if not replaced_in_run:
                                     print(f"WARNING: Found {old} in table but it is split across runs. Replacing text.")
                                     p.text = p.text.replace(old, new)
                                     count += 1

    print(f"Total replacements made: {count}")
    doc.save(template_path)
    print(f"Saved to {template_path}")

if __name__ == "__main__":
    fix_tags()
