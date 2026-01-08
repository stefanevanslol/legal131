from docx import Document
import re
import os

input_path = "backend/templates/Natale - GEICO BI Demand (3) (5).docx"
output_path = "backend/templates/Natale - GEICO BI Demand (3) (5)_fixed.docx"

def fix_template():
    if not os.path.exists(input_path):
        print(f"File not found: {input_path}")
        return

    doc = Document(input_path)
    count = 0
    
    # Iterate over all paragraphs
    # We look for paragraphs that contain "{{" and "}}"
    # If found, we replace the runs with a single run containing the full text
    # This "heals" split tags.
    
    sections_to_check = []
    
    # 1. Document Paragraphs
    sections_to_check.extend(doc.paragraphs)
    
    # 2. Table Cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                sections_to_check.extend(cell.paragraphs)
                
    for p in sections_to_check:
        text = p.text
        if "{{" in text and "}}" in text:
            # Check if tags are involved
            matches = re.findall(r"\{\{.*?\}\}", text)
            if matches:
                 # Found tags. Let's fix this paragraph.
                 print(f"Fixing paragraph: '{text[:50]}...'")
                 
                 # Store the alignment and style (basic attempt to preserve)
                 alignment = p.alignment
                 style = p.style
                 
                 # Clear existing runs
                 p.clear()
                 
                 # Add new run with full text
                 # Note: This removes character-level formatting (bold/underline) within the paragraph.
                 # But since the goal is to REMOVE bold/underline from these specific variables, this is actually a FEATURE.
                 run = p.add_run(text)
                 
                 # Restore paragraph properties if needed, but p.clear() usually keeps paragraph-level formatting.
                 # run properties (bold/italic) are gone, effectively resetting correctly.
                 
                 count += 1

    doc.save(output_path)
    print(f"Fixed {count} paragraphs.")
    print(f"Saved to {output_path}")

if __name__ == "__main__":
    fix_template()
