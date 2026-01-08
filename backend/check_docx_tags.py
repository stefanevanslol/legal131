from docx import Document
import re
import os

# Assuming this is the latest fixed template based on the name
template_path = r"backend/templates/Natale - GEICO BI Demand (3) (6).docx"

def check_tags():
    if not os.path.exists(template_path):
        print(f"File not found: {template_path}")
        return

    doc = Document(template_path)
    full_text = []
    
    # Paragraphs
    for p in doc.paragraphs:
        full_text.append(p.text)
        
    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    full_text.append(p.text)
                    
    text = "\n".join(full_text)
    
    # Regex for {{ ... }} including possible 'r ' prefix
    matches = re.findall(r"\{\{(.*?)\}\}", text)
    
    print(f"--- Analyzing {template_path} ---")
    print(f"Total Matches: {len(matches)}")
    
    unique_tags = sorted(list(set([m.strip() for m in matches])))
    
    print("\n--- Unique Tags Found ---")
    relevant_keys = ["TreatmentsAndDiagnosses", "TreatmentsAndDiagnosses2", "complaintsByclient"]
    for tag in unique_tags:
        # Check if it relates to our keys
        is_relevant = any(key in tag for key in relevant_keys)
        marker = " [RELEVANT]" if is_relevant else ""
        print(f"'{tag}'{marker}")

if __name__ == "__main__":
    check_tags()
